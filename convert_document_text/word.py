from pathlib import Path
from typing import Dict, Any, List
from docx import Document
import unicodedata


class Word:
    """
    Clase para extraer texto de documentos Word (.docx)
    """

    def __init__(self):
        print("Inicializando extractor Word")

    def _normalize_text(self, text):
        """Normaliza el texto eliminando acentos y convirtiendo a minúsculas."""
        text = text.lower()
        text = ''.join(
            c for c in unicodedata.normalize('NFD', text)
            if unicodedata.category(c) != 'Mn'
        )
        return text

    def _extraer_metadatos_word(self, doc: Document) -> Dict[str, Any]:
        """Extrae metadatos del documento Word."""
        core_props = doc.core_properties
        return {
            'titulo': core_props.title or '',
            'autor': core_props.author or '',
            'fecha_creacion': str(core_props.created) if core_props.created else '',
            'fecha_modificacion': str(core_props.modified) if core_props.modified else '',
            'parrafos': len(doc.paragraphs)
        }

    def _extraer_tablas(self, doc: Document) -> List[Dict[str, Any]]:
        """Extrae tablas del documento Word."""
        tablas_procesadas = []

        for i, tabla in enumerate(doc.tables):
            try:
                datos = []
                for fila in tabla.rows:
                    fila_datos = [celda.text.strip() for celda in fila.cells]
                    if any(fila_datos):  # Solo agregar si hay contenido
                        datos.append(fila_datos)

                if datos:
                    tablas_procesadas.append({
                        'indice': i + 1,
                        'datos': datos,
                        'filas': len(datos),
                        'columnas': len(datos[0]) if datos else 0
                    })
            except Exception as e:
                print(f"Error procesando tabla {i + 1}: {e}")
                continue

        return tablas_procesadas

    def procesar_word(self, ruta_archivo: str) -> Dict[str, Any]:
        """
        Procesa un archivo Word extrayendo texto, tablas y metadatos.

        Args:
            ruta_archivo: Ruta al archivo Word

        Returns:
            Diccionario con contenido estructurado del documento
        """
        resultado = {
            'tipo_archivo': 'word',
            'archivo': Path(ruta_archivo).name,
            'metadatos': {},
            'contenido': {
                'texto': '',
                'tablas': []
            }
        }

        # Abrir el documento Word
        doc = Document(ruta_archivo)

        try:
            # Extraer metadatos
            resultado['metadatos'] = self._extraer_metadatos_word(doc)

            # Extraer texto de párrafos
            texto_completo = []
            for parrafo in doc.paragraphs:
                texto = parrafo.text.strip()
                if texto:
                    texto_completo.append(texto)

            resultado['contenido']['texto'] = self._normalize_text(' '.join(texto_completo))

            # Extraer tablas
            resultado['contenido']['tablas'] = self._extraer_tablas(doc)

        except Exception as e:
            print(f"Error procesando documento Word: {e}")
            raise

        return resultado
