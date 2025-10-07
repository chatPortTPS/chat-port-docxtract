"""
Document Extractor - Extractor de Documentos
============================================

Módulo principal para extraer contenido de documentos PDF, Word y Excel.
Incluye funcionalidades de limpieza y filtrado de contenido irrelevante.

Autor: Assistant
Fecha: Septiembre 2025
"""

import os
import re
import json
from pathlib import Path
from datetime import datetime
from typing import Dict, List, Any, Optional, Union
import warnings

# Suprimir advertencias de bibliotecas
warnings.filterwarnings("ignore", category=UserWarning)

# Importaciones de las bibliotecas de procesamiento
try:
    import fitz  # PyMuPDF
except ImportError:
    print("Error: PyMuPDF no está instalado. Ejecute: pip install pymupdf")

try:
    import camelot
except ImportError:
    print("Error: Camelot no está instalado. Ejecute: pip install camelot-py[cv]")

try:
    from docx import Document
    from docx.table import Table
except ImportError:
    print("Error: python-docx no está instalado. Ejecute: pip install python-docx")

try:
    from openpyxl import load_workbook
except ImportError:
    print("Error: openpyxl no está instalado. Ejecute: pip install openpyxl")

import pandas as pd


class DocumentExtractor:
    """
    Extractor principal de documentos PDF, Word y Excel.
    
    Maneja la extracción de texto, tablas y metadatos con filtrado
    inteligente de contenido irrelevante.
    """
    
    def __init__(self):
        """Inicializa el extractor con configuración por defecto."""
        self.tipos_soportados = ['.pdf', '.docx', '.xlsx']
        
        # Patrones para detectar contenido a filtrar
        self.patrones_indice = [
            r'\b\d+\s*$',  # Números de página al final de línea
            r'\.{3,}\s*\d+\s*$',  # Puntos seguidos de números de página
            r'página\s+\d+',  # Palabra página seguida de número
            r'cap[íi]tulo\s+\d+.*\d+$',  # Capítulo seguido de número de página
        ]
        
        self.palabras_portada = [
            'índice', 'index', 'contenido', 'contents', 'tabla de contenido',
            'table of contents', 'resumen ejecutivo', 'executive summary'
        ]
    
    def extraer_documento(self, ruta_archivo: str) -> Dict[str, Any]:
        """
        Extrae contenido de un documento según su tipo.
        
        Args:
            ruta_archivo: Ruta al archivo a procesar
            
        Returns:
            Diccionario con el contenido extraído estructurado
            
        Raises:
            FileNotFoundError: Si el archivo no existe
            ValueError: Si el tipo de archivo no es soportado
        """
        ruta = Path(ruta_archivo)
        
        if not ruta.exists():
            raise FileNotFoundError(f"El archivo no existe: {ruta_archivo}")
        
        extension = ruta.suffix.lower()
        
        if extension not in self.tipos_soportados:
            raise ValueError(f"Tipo de archivo no soportado: {extension}")
         
        if extension == '.pdf':
            return self._procesar_pdf(str(ruta))
        elif extension == '.docx':
            return self._procesar_word(str(ruta))
        elif extension == '.xlsx':
            return self._procesar_excel(str(ruta))
    
    def _procesar_pdf(self, ruta_archivo: str) -> Dict[str, Any]:
        """
        Procesa un archivo PDF extrayendo texto, tablas y metadatos.
        
        Args:
            ruta_archivo: Ruta al archivo PDF
            
        Returns:
            Diccionario con contenido estructurado del PDF
        """
        resultado = {
            'tipo_archivo': 'pdf',
            'archivo': Path(ruta_archivo).name,
            'metadatos': {},
            'contenido': {
                'texto': '',
                'tablas': []
            }
        }
        
        # Abrir el documento PDF
        doc = fitz.open(ruta_archivo)
        
        try:
            # Extraer metadatos
            resultado['metadatos'] = self._extraer_metadatos_pdf(doc)
            
            # Extraer texto de todas las páginas
            texto_paginas = []
            cabeceras_comunes = []
            pies_comunes = []
            
            for num_pagina in range(len(doc)):
                pagina = doc[num_pagina]
                texto_pagina = pagina.get_text()
                
                if texto_pagina.strip():
                    texto_paginas.append(texto_pagina)
                    
                    # Detectar posibles cabeceras y pies
                    lineas = texto_pagina.strip().split('\n')
                    if lineas:
                        cabeceras_comunes.append(lineas[0])
                        if len(lineas) > 1:
                            pies_comunes.append(lineas[-1])
            
            # Identificar cabeceras y pies repetitivos
            cabeceras_repetitivas = self._encontrar_texto_repetitivo(cabeceras_comunes)
            pies_repetitivos = self._encontrar_texto_repetitivo(pies_comunes)
            
            # Limpiar y procesar el texto
            texto_completo = self._limpiar_texto_pdf(
                texto_paginas, cabeceras_repetitivas, pies_repetitivos
            )
            
            resultado['contenido']['texto'] = texto_completo
            
            # Extraer tablas usando Camelot
            try:
                tablas = camelot.read_pdf(ruta_archivo, pages='all', flavor='lattice')
                resultado['contenido']['tablas'].extend(
                    self._procesar_tablas_camelot(tablas)
                )
            except Exception as e:
                print(f"Advertencia: No se pudieron extraer tablas con Camelot: {e}")
                # Intentar con flavor 'stream' como alternativa
                try:
                    tablas = camelot.read_pdf(ruta_archivo, pages='all', flavor='stream')
                    resultado['contenido']['tablas'].extend(
                        self._procesar_tablas_camelot(tablas)
                    )
                except Exception as e2:
                    print(f"Advertencia: Falló también la extracción alternativa de tablas: {e2}")
        
        finally:
            doc.close()
        
        return resultado
    
    def _procesar_word(self, ruta_archivo: str) -> Dict[str, Any]:
        """
        Procesa un archivo Word extrayendo texto, tablas y metadatos.
        
        Args:
            ruta_archivo: Ruta al archivo Word
            
        Returns:
            Diccionario con contenido estructurado del Word
        """
        resultado = {
            'tipo_archivo': 'docx',
            'archivo': Path(ruta_archivo).name,
            'metadatos': {},
            'contenido': {
                'texto': '',
                'tablas': []
            }
        }
        
        # Abrir el documento Word
        doc = Document(ruta_archivo)
        
        # Extraer metadatos
        resultado['metadatos'] = self._extraer_metadatos_word(doc)
        
        # Extraer texto de párrafos
        texto_parrafos = []
        tablas_encontradas = []
        
        for elemento in doc.element.body:
            if elemento.tag.endswith('p'):  # Párrafo
                parrafo = next((p for p in doc.paragraphs if p._element == elemento), None)
                if parrafo and parrafo.text.strip():
                    texto_parrafos.append(parrafo.text.strip())
            
            elif elemento.tag.endswith('tbl'):  # Tabla
                tabla = next((t for t in doc.tables if t._element == elemento), None)
                if tabla:
                    tablas_encontradas.append(tabla)
        
        # Procesar texto limpiando contenido irrelevante
        texto_limpio = self._limpiar_texto_word(texto_parrafos)
        resultado['contenido']['texto'] = texto_limpio
        
        # Procesar tablas
        resultado['contenido']['tablas'] = self._procesar_tablas_word(tablas_encontradas)
        
        return resultado
    
    def _procesar_excel(self, ruta_archivo: str) -> Dict[str, Any]:
        """
        Procesa un archivo Excel extrayendo datos de todas las hojas y metadatos.
        
        Args:
            ruta_archivo: Ruta al archivo Excel
            
        Returns:
            Diccionario con contenido estructurado del Excel
        """
        resultado = {
            'tipo_archivo': 'xlsx',
            'archivo': Path(ruta_archivo).name,
            'metadatos': {},
            'contenido': {
                'texto': '',
                'tablas': []
            }
        }
        
        # Abrir el libro de Excel
        wb = load_workbook(ruta_archivo, read_only=True, data_only=True)
        
        try:
            # Extraer metadatos
            resultado['metadatos'] = self._extraer_metadatos_excel(wb)
            
            # Procesar todas las hojas
            for i, nombre_hoja in enumerate(wb.sheetnames):
                hoja = wb[nombre_hoja]
                datos_hoja = self._procesar_hoja_excel(hoja, nombre_hoja, i)
                
                if datos_hoja:
                    resultado['contenido']['tablas'].append(datos_hoja)
            
            # El texto para Excel será un resumen de las hojas
            resultado['contenido']['texto'] = self._generar_resumen_excel(resultado['contenido']['tablas'])
        
        finally:
            wb.close()
        
        return resultado
    
    def _extraer_metadatos_pdf(self, doc: fitz.Document) -> Dict[str, Any]:
        """Extrae metadatos de un documento PDF."""
        metadatos = doc.metadata
        return {
            'titulo': metadatos.get('title', ''),
            'autor': metadatos.get('author', ''),
            'fecha_creacion': metadatos.get('creationDate', ''),
            'fecha_modificacion': metadatos.get('modDate', ''),
            'paginas': len(doc)
        }
    
    def _extraer_metadatos_word(self, doc: Document) -> Dict[str, Any]:
        """Extrae metadatos de un documento Word."""
        props = doc.core_properties
        return {
            'titulo': props.title or '',
            'autor': props.author or '',
            'fecha_creacion': props.created.isoformat() if props.created else '',
            'fecha_modificacion': props.modified.isoformat() if props.modified else '',
            'parrafos': len(doc.paragraphs)
        }
    
    def _extraer_metadatos_excel(self, wb) -> Dict[str, Any]:
        """Extrae metadatos de un libro Excel."""
        props = wb.properties
        return {
            'titulo': props.title or '',
            'autor': props.creator or '',
            'fecha_creacion': props.created.isoformat() if props.created else '',
            'fecha_modificacion': props.modified.isoformat() if props.modified else '',
            'hojas': len(wb.sheetnames)
        }
    
    def _encontrar_texto_repetitivo(self, lineas: List[str]) -> List[str]:
        """Encuentra texto que se repite en múltiples páginas."""
        if len(lineas) < 2:
            return []
        
        contador = {}
        for linea in lineas:
            linea = linea.strip()
            if linea and len(linea) > 5:  # Ignorar líneas muy cortas
                contador[linea] = contador.get(linea, 0) + 1
        
        # Considerar repetitivo si aparece en más del 50% de las páginas
        umbral = len(lineas) * 0.5
        repetitivos = [linea for linea, count in contador.items() if count > umbral]
        
        return repetitivos
    
    def _limpiar_texto_pdf(self, texto_paginas: List[str], 
                          cabeceras_repetitivas: List[str], 
                          pies_repetitivos: List[str]) -> str:
        """Limpia el texto del PDF removiendo contenido irrelevante."""
        texto_limpio = []
        
        for i, texto_pagina in enumerate(texto_paginas):
            lineas = texto_pagina.strip().split('\n')
            lineas_filtradas = []
            
            for linea in lineas:
                linea = linea.strip()
                
                # Saltar líneas vacías
                if not linea:
                    continue
                
                # Filtrar cabeceras y pies repetitivos
                if linea in cabeceras_repetitivas or linea in pies_repetitivos:
                    continue
                
                # Filtrar líneas que parecen índice
                if self._es_linea_indice(linea):
                    continue
                
                # Si es la primera página y parece portada, saltar
                if i == 0 and self._es_posible_portada(linea):
                    continue
                
                lineas_filtradas.append(linea)
            
            # Reconstruir párrafos
            if lineas_filtradas:
                texto_reconstruido = self._reconstruir_parrafos(lineas_filtradas)
                texto_limpio.append(texto_reconstruido)
        
        return '\n\n'.join(texto_limpio)
    
    def _limpiar_texto_word(self, parrafos: List[str]) -> str:
        """Limpia el texto de Word removiendo contenido irrelevante."""
        parrafos_filtrados = []
        
        for i, parrafo in enumerate(parrafos):
            # Filtrar líneas que parecen índice
            if self._es_linea_indice(parrafo):
                continue
            
            # Si es el primer párrafo y parece portada, evaluar
            if i == 0 and self._es_posible_portada(parrafo):
                continue
            
            parrafos_filtrados.append(parrafo)
        
        return '\n\n'.join(parrafos_filtrados)
    
    def _es_linea_indice(self, linea: str) -> bool:
        """Determina si una línea parece ser parte de un índice."""
        for patron in self.patrones_indice:
            if re.search(patron, linea, re.IGNORECASE):
                return True
        return False
    
    def _es_posible_portada(self, linea: str) -> bool:
        """Determina si una línea podría ser parte de una portada."""
        linea_lower = linea.lower()
        return any(palabra in linea_lower for palabra in self.palabras_portada)
    
    def _reconstruir_parrafos(self, lineas: List[str]) -> str:
        """Reconstruye párrafos a partir de líneas fragmentadas."""
        parrafos = []
        parrafo_actual = []
        
        for linea in lineas:
            # Si la línea termina con punto, es fin de párrafo
            if linea.endswith('.') or linea.endswith('!') or linea.endswith('?'):
                parrafo_actual.append(linea)
                parrafos.append(' '.join(parrafo_actual))
                parrafo_actual = []
            # Si la línea parece ser continuación
            elif parrafo_actual and linea[0].islower():
                parrafo_actual.append(linea)
            # Nueva oración/párrafo
            else:
                if parrafo_actual:
                    parrafos.append(' '.join(parrafo_actual))
                parrafo_actual = [linea]
        
        # Agregar el último párrafo si existe
        if parrafo_actual:
            parrafos.append(' '.join(parrafo_actual))
        
        return '\n\n'.join(parrafos)
    
    def _procesar_tablas_camelot(self, tablas) -> List[Dict[str, Any]]:
        """Procesa tablas extraídas con Camelot."""
        tablas_procesadas = []
        
        for i, tabla in enumerate(tablas):
            try:
                df = tabla.df
                
                # Convertir a lista de diccionarios
                if not df.empty:
                    # Usar la primera fila como encabezados si no están vacíos
                    if df.iloc[0].notna().all():
                        df.columns = df.iloc[0]
                        df = df.drop(df.index[0])
                    
                    # Filtrar filas completamente vacías
                    df = df.dropna(how='all')
                    
                    if not df.empty:
                        datos = df.to_dict('records')
                        tablas_procesadas.append({
                            'indice': i + 1,
                            'datos': datos,
                            'filas': len(datos),
                            'columnas': len(df.columns)
                        })
            
            except Exception as e:
                print(f"Error procesando tabla {i + 1}: {e}")
                continue
        
        return tablas_procesadas
    
    def _procesar_tablas_word(self, tablas: List[Table]) -> List[Dict[str, Any]]:
        """Procesa tablas extraídas de un documento Word."""
        tablas_procesadas = []
        
        for i, tabla in enumerate(tablas):
            try:
                datos = []
                
                # Obtener los datos de cada fila
                for fila in tabla.rows:
                    fila_datos = []
                    for celda in fila.cells:
                        texto_celda = celda.text.strip()
                        fila_datos.append(texto_celda)
                    datos.append(fila_datos)
                
                # Usar la primera fila como encabezados
                if datos:
                    encabezados = datos[0]
                    filas_datos = datos[1:]
                    
                    # Convertir a lista de diccionarios
                    datos_dict = []
                    for fila in filas_datos:
                        fila_dict = {}
                        for j, valor in enumerate(fila):
                            encabezado = encabezados[j] if j < len(encabezados) else f"Columna_{j+1}"
                            fila_dict[encabezado] = valor
                        datos_dict.append(fila_dict)
                    
                    tablas_procesadas.append({
                        'indice': i + 1,
                        'datos': datos_dict,
                        'filas': len(datos_dict),
                        'columnas': len(encabezados)
                    })
            
            except Exception as e:
                print(f"Error procesando tabla Word {i + 1}: {e}")
                continue
        
        return tablas_procesadas
    
    def _procesar_hoja_excel(self, hoja, nombre_hoja: str, indice: int) -> Optional[Dict[str, Any]]:
        """Procesa una hoja de Excel extrayendo sus datos."""
        try:
            datos = []
            
            # Encontrar el rango de datos (evitar celdas vacías al final)
            max_row = hoja.max_row
            max_col = hoja.max_column
            
            # Leer los datos
            for fila in range(1, max_row + 1):
                fila_datos = []
                for col in range(1, max_col + 1):
                    celda = hoja.cell(row=fila, column=col)
                    valor = celda.value
                    if valor is not None:
                        fila_datos.append(str(valor))
                    else:
                        fila_datos.append('')
                
                # Solo agregar filas que tengan al menos un valor
                if any(celda.strip() for celda in fila_datos):
                    datos.append(fila_datos)
            
            if not datos:
                return None
            
            # Usar la primera fila como encabezados
            encabezados = datos[0]
            filas_datos = datos[1:] if len(datos) > 1 else []
            
            # Convertir a lista de diccionarios
            datos_dict = []
            for fila in filas_datos:
                fila_dict = {}
                for j, valor in enumerate(fila):
                    encabezado = encabezados[j] if j < len(encabezados) else f"Columna_{j+1}"
                    fila_dict[encabezado] = valor
                datos_dict.append(fila_dict)
            
            return {
                'indice': indice + 1,
                'nombre_hoja': nombre_hoja,
                'datos': datos_dict,
                'filas': len(datos_dict),
                'columnas': len(encabezados)
            }
        
        except Exception as e:
            print(f"Error procesando hoja '{nombre_hoja}': {e}")
            return None
    
    def _generar_resumen_excel(self, tablas: List[Dict[str, Any]]) -> str:
        """Genera un resumen textual del contenido de las hojas de Excel."""
        if not tablas:
            return "Archivo Excel sin datos relevantes."
        
        resumen = []
        resumen.append(f"Archivo Excel con {len(tablas)} hoja(s) de datos:")
        
        for tabla in tablas:
            nombre = tabla.get('nombre_hoja', f"Hoja {tabla['indice']}")
            filas = tabla['filas']
            columnas = tabla['columnas']
            resumen.append(f"- {nombre}: {filas} filas × {columnas} columnas")
        
        return '\n'.join(resumen)


def main():
    """Función principal para ejecutar el extractor desde línea de comandos."""
    import sys
    
    if len(sys.argv) != 2:
        print("Uso: python document_extractor.py <ruta_archivo>")
        print("Tipos soportados: .pdf, .docx, .xlsx")
        return
    
    ruta_archivo = sys.argv[1]
    
    try:
        extractor = DocumentExtractor()
        resultado = extractor.extraer_documento(ruta_archivo)
        
        # Mostrar resultado en formato JSON con indentación
        print("\n" + "="*50)
        print("RESULTADO DE LA EXTRACCIÓN")
        print("="*50)
        print(json.dumps(resultado, indent=2, ensure_ascii=False))
    
    except Exception as e:
        print(f"Error: {e}")


if __name__ == "__main__":
    main()
